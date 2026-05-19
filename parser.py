import pdfplumber
from docx import Document
import re

def get_quizzes_programming(file_path):
    """
    Dasturlash fani uchun maxsus parser.
    To'g'ri javob har doim BIRINCHI variant (— bilan boshlanadi).
    """
    try:
        doc = Document(file_path)
        quiz_data = []
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        i = 0
        while i < len(paragraphs):
            if "Savol" in paragraphs[i] or "савол" in paragraphs[i].lower():
                question_parts = []
                i += 1
                while i < len(paragraphs) and not paragraphs[i].startswith("—"):
                    question_parts.append(paragraphs[i])
                    i += 1
                
                question_text = " ".join(question_parts).strip()
                options = []
                while i < len(paragraphs) and paragraphs[i].startswith("—"):
                    option_text = paragraphs[i].replace("—", "").strip()
                    options.append(option_text[:100])
                    i += 1
                
                if len(options) >= 2 and question_text:
                    quiz_data.append({
                        "question": question_text[:250],
                        "options": options[:10],
                        "correct": 0
                    })
            else:
                i += 1
        return quiz_data
    except Exception as e:
        print(f"Dasturlash parser xatosi: {e}")
        return []

def get_quizzes(file_path):
    """
    Falsafa, MT-V-A, Dinshunoslik va Ingliz tili-{Di} fanlari uchun standart parser.
    Format: ++++ va ==== bilan ajratilgan
    """
    try:
        doc = Document(file_path)
        full_text = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
        raw_questions = full_text.split("++++")
        quiz_data = []

        for item in raw_questions:
            item = item.strip()
            if not item: continue
            parts = item.split("====")
            if len(parts) < 2: continue
            
            question_text = parts[0].strip()
            options_raw = [p.strip() for p in parts[1:] if p.strip()]
            
            correct_index = 0
            final_options = []
            for i, opt in enumerate(options_raw):
                clean_opt = opt.replace("#", "").strip()
                if opt.startswith("#"):
                    correct_index = i
                final_options.append(clean_opt[:100])
            
            if len(final_options) >= 2:
                quiz_data.append({
                    "question": question_text[:250],
                    "options": final_options[:10],
                    "correct": correct_index
                })
        return quiz_data
    except Exception as e:
        print(f"Standart parser xatosi: {e}")
        return []

def normalize_text(text):
    """
    Matnni taqqoslash uchun normalizatsiya qiladi
    """
    if not text:
        return ""
    text = re.sub(r'[^a-zA-Z0-9\s]', '', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip().lower()

# Inglizcha savol prefixlarini uzbekchaga tarjima qilish
ENGLISH_QUESTION_TRANSLATIONS = {
    "choose the correct answer": "To'g'ri javobni tanlang",
    "choose the correct word": "To'g'ri so'zni tanlang",
    "choose the word closest in meaning to": "Ma'nosi eng yaqin so'zni tanlang",
    "choose the word closest in meaning": "Ma'nosi eng yaqin so'zni tanlang",
    "complete the sentence": "Gapni to'ldiring",
    "choose the correct alternative to": "To'g'ri muqobilni tanlang",
    "choose the correct translation of the word": "So'zning to'g'ri tarjimasini tanlang",
    "choose the correct preposition": "To'g'ri predlogni tanlang",
    "it is a group of people working together for illegal purposes": "Bu noqonuniy maqsadlarda birga ishlaydigan odamlar guruhi",
    "choose the best": "Eng yaxshisini tanlang",
    "select the correct": "To'g'risini tanlang"
}

def translate_english_question(question_text):
    """
    Inglizcha savol matnini tarjima bilan qaytaradi.
    Format: {Tarjima = Inglizcha}
    """
    question_lower = question_text.lower().strip()
    
    for eng_phrase, uz_translation in ENGLISH_QUESTION_TRANSLATIONS.items():
        if eng_phrase in question_lower:
            start_idx = question_lower.find(eng_phrase)
            end_idx = start_idx + len(eng_phrase)
            original_eng = question_text[start_idx:end_idx]
            remaining = question_text[end_idx:].strip()
            
            if remaining:
                return f"{uz_translation} = {original_eng} {remaining}"
            else:
                return f"{uz_translation} = {original_eng}"
    
    return question_text

def calculate_similarity(text1, text2):
    """
    Ikki matn o'rtasidagi o'xshashlikni hisoblash (0-100)
    """
    norm1 = normalize_text(text1)
    norm2 = normalize_text(text2)
    
    if not norm1 or not norm2:
        return 0
    
    if norm1 == norm2:
        return 100
    
    if norm1 in norm2:
        return (len(norm1) / len(norm2)) * 100
    
    if norm2 in norm1:
        return (len(norm2) / len(norm1)) * 100
    
    words1 = set(norm1.split())
    words2 = set(norm2.split())
    
    if not words1 or not words2:
        return 0
    
    intersection = len(words1.intersection(words2))
    union = len(words1.union(words2))
    
    if union == 0:
        return 0
    
    return (intersection / union) * 100

def get_quizzes_english_pdf_docx(docx_path, pdf_path):
    """
    Ingliz tili-{KIN} uchun PDF+DOCX parser.
    """
    try:
        print("\n" + "="*60)
        print("🔍 INGLIZ TILI PARSER BOSHLANDI")
        print("="*60)
        
        pdf_red_texts = []
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                chars = page.chars
                current_word = []
                
                for char in chars:
                    color = char.get("non_stroking_color") or char.get("stroking_color")
                    is_red = False
                    
                    if color and len(color) == 3:
                        r, g, b = color
                        if (r > 0.6 and g < 0.4 and b < 0.4) or \
                           (isinstance(r, int) and r > 150 and g < 100 and b < 100):
                            is_red = True
                    
                    if is_red:
                        current_word.append(char["text"])
                    else:
                        if current_word:
                            word_str = "".join(current_word).strip()
                            if word_str and not word_str.isdigit() and len(word_str) > 1:
                                pdf_red_texts.append(word_str)
                            current_word = []
                
                if current_word:
                    word_str = "".join(current_word).strip()
                    if word_str and not word_str.isdigit() and len(word_str) > 1:
                        pdf_red_texts.append(word_str)
        
        doc = Document(docx_path)
        raw_lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        clean_lines = []
        for line in raw_lines:
            if set(line).issubset({'_', '=', '+', '-', '*', ' ', '—'}) and len(line) > 1:
                continue
            clean_lines.append(line)

        i = 0
        docx_quizzes = []
        
        while i < len(clean_lines):
            line = clean_lines[i]
            if "choose the correct" in line.lower() or "choose the word" in line.lower() or \
               "choose the best" in line.lower() or "select the correct" in line.lower() or \
               "complete the sentence" in line.lower() or "it is a group" in line.lower():
                question_text = line
                options = []
                i += 1
                
                while i < len(clean_lines):
                    current_line = clean_lines[i]
                    if "choose the correct" in current_line.lower() or \
                       "choose the word" in current_line.lower() or \
                       "choose the best" in current_line.lower() or \
                       "select the correct" in current_line.lower() or \
                       "complete the sentence" in current_line.lower() or \
                       "it is a group" in current_line.lower():
                        break
                    
                    if current_line and len(current_line) < 200:
                        options.append(current_line)
                    i += 1
                
                if len(options) >= 2:
                    translated_question = translate_english_question(question_text)
                    docx_quizzes.append({
                        "question": translated_question,
                        "options": options,
                        "original_question": question_text
                    })
            else:
                i += 1
        
        quiz_data = []
        for quiz_idx, docx_quiz in enumerate(docx_quizzes):
            best_correct_idx = 0
            best_pdf_answer = ""
            best_overall_score = 0
            
            for pdf_text in pdf_red_texts:
                for opt_idx, option in enumerate(docx_quiz["options"]):
                    similarity = calculate_similarity(pdf_text, option)
                    if similarity > best_overall_score:
                        best_overall_score = similarity
                        best_correct_idx = opt_idx
                        best_pdf_answer = pdf_text
            
            quiz_data.append({
                "question": docx_quiz["question"][:300],
                "options": [opt[:100] for opt in docx_quiz["options"]][:10],
                "correct": best_correct_idx,
                "pdf_answer": best_pdf_answer if best_pdf_answer else "Topilmadi",
                "confidence": best_overall_score
            })
        
        return quiz_data
    except Exception as e:
        print(f"❌ Ingliz tili PDF+DOCX parser xatosi: {e}")
        return []
